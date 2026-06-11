"""
SRR Billing Generator — Generator Kwitansi & Surat Tagihan KJPP SRR
v1.0 | Streamlit Cloud App
"""

import streamlit as st
import anthropic
import json
import io
import zipfile
import re
from datetime import datetime, date

# ── python-docx (template docx)
from docx import Document
from docx.shared import Pt

import base64

# ─────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────
st.set_page_config(
    page_title="SRR Billing Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────
# CUSTOM CSS  — clean professional look
# ─────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

html, body, [class*="css"] { font-family: 'Plus Jakarta Sans', sans-serif; }

.srr-header {
    background: linear-gradient(135deg, #0f3460 0%, #16213e 100%);
    color: white;
    padding: 1.5rem 2rem;
    border-radius: 12px;
    margin-bottom: 1.5rem;
    display: flex;
    align-items: center;
    gap: 1rem;
}
.srr-header h1 { margin:0; font-size:1.6rem; font-weight:700; letter-spacing:-0.5px; }
.srr-header p  { margin:0; font-size:0.85rem; opacity:0.75; }

.step-pill {
    display: inline-flex; align-items: center; gap: 6px;
    background: #eef2ff; color: #3730a3;
    border-radius: 20px; padding: 4px 12px;
    font-size: 0.78rem; font-weight: 600;
    margin-bottom: 0.5rem;
}

.project-card {
    border: 1.5px solid #e2e8f0;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.8rem;
    background: white;
    transition: box-shadow 0.2s;
}
.project-card:hover { box-shadow: 0 4px 16px rgba(0,0,0,0.08); }

.tag-pill {
    display: inline-block;
    background: #f0fdf4; color: #166534;
    border-radius: 6px; padding: 2px 8px;
    font-size: 0.75rem; font-weight: 600;
    font-family: 'JetBrains Mono', monospace;
}

.amount-display {
    font-family: 'JetBrains Mono', monospace;
    font-weight: 600;
    color: #0f3460;
    font-size: 1.1rem;
}

.stButton > button {
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-family: 'Plus Jakarta Sans', sans-serif !important;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# TERBILANG RUPIAH
# ─────────────────────────────────────────
def terbilang_rupiah(n: int) -> str:
    satuan = ['', 'Satu', 'Dua', 'Tiga', 'Empat', 'Lima', 'Enam', 'Tujuh', 'Delapan', 'Sembilan',
              'Sepuluh', 'Sebelas', 'Dua Belas', 'Tiga Belas', 'Empat Belas', 'Lima Belas',
              'Enam Belas', 'Tujuh Belas', 'Delapan Belas', 'Sembilan Belas']
    puluhan = ['', '', 'Dua Puluh', 'Tiga Puluh', 'Empat Puluh', 'Lima Puluh',
               'Enam Puluh', 'Tujuh Puluh', 'Delapan Puluh', 'Sembilan Puluh']

    def _t(n):
        if n == 0:
            return ''
        elif n < 20:
            return satuan[n]
        elif n < 100:
            ones = satuan[n % 10]
            return (puluhan[n // 10] + (' ' + ones if ones else '')).strip()
        elif n < 200:
            rest = _t(n - 100)
            return 'Seratus' + (' ' + rest if rest else '')
        elif n < 1000:
            rest = _t(n % 100)
            return satuan[n // 100] + ' Ratus' + (' ' + rest if rest else '')
        elif n < 2000:
            rest = _t(n - 1000)
            return 'Seribu' + (' ' + rest if rest else '')
        elif n < 1_000_000:
            rest = _t(n % 1000)
            return _t(n // 1000) + ' Ribu' + (' ' + rest if rest else '')
        elif n < 1_000_000_000:
            rest = _t(n % 1_000_000)
            return _t(n // 1_000_000) + ' Juta' + (' ' + rest if rest else '')
        elif n < 1_000_000_000_000:
            rest = _t(n % 1_000_000_000)
            return _t(n // 1_000_000_000) + ' Miliar' + (' ' + rest if rest else '')
        else:
            rest = _t(n % 1_000_000_000_000)
            return _t(n // 1_000_000_000_000) + ' Triliun' + (' ' + rest if rest else '')

    if n == 0:
        return 'Nol Rupiah'
    return _t(int(n)) + ' Rupiah'


# ─────────────────────────────────────────
# FEE CALCULATOR
# ─────────────────────────────────────────
def hitung_tagihan(imbalan_jasa: float, tarif_ppn: float = 0.12) -> dict:
    """Hitung DPP + PPN + Total sesuai pola KJPP SRR (DPP = IJ × 11/12)"""
    ij = round(imbalan_jasa)
    dpb = round(ij * 11 / 12)
    ppn = round(dpb * tarif_ppn)
    total = ij + ppn
    return {
        "imbalan_jasa": ij,
        "dpb_ppn": dpb,
        "ppn": ppn,
        "total": total,
        "terbilang": terbilang_rupiah(total),
    }


# ─────────────────────────────────────────
# NOMOR SURAT GENERATOR
# ─────────────────────────────────────────
def generate_nomor(tanggal: date, seq: int, tipe: str, kode_klien: str) -> str:
    """
    Format: YYMMDD.SEQ/SRR-JK/TIPE/KODEKLIEN
    Contoh KWT : 260115.001/SRR-JK/KWT.PJK/IDXSTI
    Contoh SK-OR: 260115.001/SRR-JK/SK-OR/IDXSTI
    """
    prefix = tanggal.strftime("%y%m%d")
    seq_str = f"{seq:03d}"
    kode = kode_klien.upper().replace(" ", "").replace(".", "")[:8]
    return f"{prefix}.{seq_str}/SRR-JK/{tipe}/{kode}"


def format_tanggal_indo(d: date) -> str:
    bulan = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
             "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    return f"{d.day} {bulan[d.month]} {d.year}"


# ─────────────────────────────────────────
# CLAUDE API — EXTRACT FROM PDF/DOCX
# ─────────────────────────────────────────
_EXTRACT_PROMPT = """Kamu adalah asisten ekstraksi data dari proposal jasa penilaian KJPP SRR (Suwendho Rinaldy dan Rekan).

Ekstrak informasi berikut dan kembalikan HANYA JSON valid (tanpa komentar, tanpa markdown backtick).

PETUNJUK PENTING untuk imbalan_jasa_total:
- Cari bagian "Imbalan Jasa", "Fee", atau "Biaya" dalam proposal
- Ambil TOTAL sebelum PPN (bukan nilai setelah ditambah PPN)
- Jika ada beberapa komponen fee, JUMLAHKAN semuanya
- Nilai ini biasanya ada di baris "Jumlah Imbalan Jasa/Fee" atau "Total Fee"
- Contoh: "Rp 200.000.000" → 200000000
- Jika benar-benar tidak ditemukan, isi 0

Field JSON yang dibutuhkan:
{
  "nama_klien": "nama lengkap perusahaan klien (PT/CV/instansi)",
  "nama_klien_singkat": "akronim atau kode klien 2-8 huruf, contoh: PTRO, IDXSTI, BPID",
  "alamat_baris1": "baris pertama alamat klien (gedung/nama lokasi)",
  "alamat_baris2": "baris kedua alamat klien (nama jalan)",
  "kota": "nama kota (contoh: Jakarta, Jakarta Selatan)",
  "kode_pos": "kode pos 5 digit sebagai string",
  "up": "jabatan penerima surat (contoh: Direksi, Direktur Utama)",
  "jenis_pekerjaan": "deskripsi singkat jenis pekerjaan tanpa nama klien",
  "nomor_proposal": "nomor proposal lengkap",
  "tanggal_proposal": "tanggal proposal format DD Bulan YYYY (contoh: 12 Januari 2026)",
  "imbalan_jasa_total": 0,
  "receiver": "nama penandatangan dari SRR di akhir proposal (biasanya Ocky Rinaldy)"
}

Jika field tidak ditemukan: string kosong "" untuk teks, 0 untuk angka.

JSON:"""


def extract_proposal_data(file_bytes: bytes, filename: str) -> dict:
    """Use Claude API to extract structured data from proposal PDF/DOCX.
    PDF dikirim langsung ke API sebagai dokumen; DOCX diekstrak teksnya.
    """
    client = anthropic.Anthropic()

    if filename.lower().endswith(".pdf"):
        # Kirim PDF langsung ke Claude — tidak perlu library PDF parsing
        pdf_b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_b64,
                        },
                    },
                    {"type": "text", "text": _EXTRACT_PROMPT},
                ],
            }],
        )
    else:
        # DOCX: ekstrak teks lalu kirim ke Claude
        doc = Document(io.BytesIO(file_bytes))
        raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:8000]
        prompt = _EXTRACT_PROMPT.replace("JSON:", f"TEKS PROPOSAL:\n{raw_text}\n\nJSON:")
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
        )

    raw = response.content[0].text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        # Fallback: return empty template
        data = {
            "nama_klien": "", "nama_klien_singkat": "",
            "alamat_baris1": "", "alamat_baris2": "",
            "kota": "", "kode_pos": "", "up": "Direksi",
            "jenis_pekerjaan": "", "nomor_proposal": "",
            "tanggal_proposal": "", "imbalan_jasa_total": 0,
            "tanggal_tagihan": "", "receiver": "Ocky Rinaldy",
            "tagih_ke": "Pertama",
        }

    # Ensure tagih_ke has default value
    if not data.get("tagih_ke"):
        data["tagih_ke"] = "Pertama"

    # Ensure numeric
    try:
        data["imbalan_jasa_total"] = int(str(data.get("imbalan_jasa_total", 0)).replace(",", "").replace(".", ""))
    except (ValueError, TypeError):
        data["imbalan_jasa_total"] = 0

    return data


# ─────────────────────────────────────────
# DOKUMEN GENERATOR (python-docx) — Surat Tagihan + Kwitansi
# ─────────────────────────────────────────
def _normalize_para(para):
    """Rebuild paragraph full text from all runs, return joined string."""
    return "".join(r.text for r in para.runs)

def _rewrite_para(para, new_text):
    """
    Write new_text into paragraph preserving formatting of first run.
    Clears all other runs.
    """
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""

def replace_in_paragraph(para, replacements):
    """
    Replace placeholders in a paragraph.
    Handles two cases:
      1. {{PLACEHOLDER}} style — exact token match
      2. Heuristic replace — teks lama diganti teks baru
    Both handle placeholders split across multiple runs.
    """
    full_text = _normalize_para(para)
    changed = False
    for k, v in replacements.items():
        if k in full_text:
            full_text = full_text.replace(k, str(v))
            changed = True
    if changed:
        _rewrite_para(para, full_text)

def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, replacements)

def _detect_template_mode(template_bytes: bytes) -> str:
    """
    Detects whether a DOCX template uses {{PLACEHOLDER}} style or is a filled document.
    Returns 'placeholder' or 'heuristic'.
    """
    doc = Document(io.BytesIO(template_bytes))
    full = " ".join(p.text for p in doc.paragraphs)
    if "{{" in full and "}}" in full:
        return "placeholder"
    return "heuristic"

def _build_heuristic_replacements(template_bytes: bytes, project: dict,
                                   nomor_sk: str, tanggal_str: str, fin: dict) -> dict:
    """
    Smart heuristic replace untuk template DOCX SRR.
    Mendukung:
      a) Template x-placeholder (xxxxxx.0xx/SRR-JK/SK-OR/KODE)
      b) Template berisi surat lama (dipakai ulang dengan ganti data)
    """
    import re
    doc = Document(io.BytesIO(template_bytes))
    full_text = " ".join(p.text for p in doc.paragraphs)
    replacements = {}

    bulan_list = ["Januari","Februari","Maret","April","Mei","Juni",
                  "Juli","Agustus","September","Oktober","November","Desember"]

    # ── 1. Nomor surat: x-placeholder (xxxxxx.0xx) atau format lama (YYMMDD.NNN)
    x_nomor = re.search(r'x+[x\.\d]*/SRR-JK/SK-OR/\S+', full_text)
    real_nomor = re.search(r'\d{6}\.\d{3}/SRR-JK/SK-OR/\S+', full_text)
    if x_nomor:
        replacements[x_nomor.group(0)] = nomor_sk
    elif real_nomor:
        replacements[real_nomor.group(0)] = nomor_sk

    # ── 2. Tanggal surat (x-placeholder "xx Bulan YYYY" atau tanggal nyata)
    tgl_pattern = r'\d{1,2}\s+(?:' + '|'.join(bulan_list) + r')\s+\d{4}'
    x_tgl = re.search(r'xx\s+(?:' + '|'.join(bulan_list) + r')\s+\d{4}', full_text)
    tanggal_lama_all = re.findall(tgl_pattern, full_text)
    if x_tgl:
        replacements[x_tgl.group(0)] = tanggal_str
    elif tanggal_lama_all:
        replacements[tanggal_lama_all[0]] = tanggal_str

    # ── 3. Nama klien (baris setelah "Kepada Yth.")
    lines_txt = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    try:
        idx_kepada = next(i for i, l in enumerate(lines_txt) if "Kepada Yth" in l)
        if idx_kepada + 1 < len(lines_txt):
            nama_lama = lines_txt[idx_kepada + 1].rstrip()
            if nama_lama and nama_lama != project["nama_klien"]:
                replacements[nama_lama] = project["nama_klien"]
    except StopIteration:
        pass

    # ── 4. Alamat klien (baris gedung/jalan)
    addr_kw = ["Jl.", "Jalan", "Wisma", "Gedung", "Tower", "Lt.", "Kav", "Plaza", "Office"]
    addr_paras = [p.text.rstrip() for p in doc.paragraphs
                  if p.text.strip() and any(k in p.text for k in addr_kw)]
    if len(addr_paras) >= 1 and addr_paras[0] != project["alamat_baris1"]:
        replacements[addr_paras[0]] = project["alamat_baris1"]
    if len(addr_paras) >= 2 and addr_paras[1] != project["alamat_baris2"]:
        replacements[addr_paras[1]] = project["alamat_baris2"]

    # ── 5. Kota + kode pos
    kota_re = re.search(
        r'(?:Jakarta\s*(?:Selatan|Pusat|Utara|Barat|Timur)?|Surabaya|Bandung|Tangerang\s*Selatan)\s+\d{4,5}',
        full_text)
    if kota_re:
        old_kp = kota_re.group(0)
        new_kp = f"{project['kota']} {project['kode_pos']}".strip()
        if old_kp != new_kp:
            replacements[old_kp] = new_kp

    # ── 6. Nomor proposal
    prop_re = re.search(r'\d{6}[\d\.]+/SRR-JK/SPN[-\w]*/\S+', full_text)
    if prop_re:
        old_prop = prop_re.group(0)
        if old_prop != project["nomor_proposal"]:
            replacements[old_prop] = project["nomor_proposal"]

    # ── 7. Tanggal proposal (tanggal ke-2, termasuk non-breaking space)
    tgl_nbsp = re.search(r'\d{1,2}\xa0(?:' + '|'.join(bulan_list) + r')\s+\d{4}', full_text)
    if tgl_nbsp and tgl_nbsp.group(0) != project["tanggal_proposal"]:
        replacements[tgl_nbsp.group(0)] = project["tanggal_proposal"]
    elif len(tanggal_lama_all) >= 2 and tanggal_lama_all[1] != project["tanggal_proposal"]:
        replacements[tanggal_lama_all[1]] = project["tanggal_proposal"]

    # ── 8. Terbilang (diapit kurung)
    terb_re = re.search(r'\(\s+([A-Z][a-zA-Z\s]+Rupiah)\s+\)', full_text)
    if terb_re:
        replacements[terb_re.group(1)] = fin["terbilang"]

    # ── 9. Jenis pekerjaan di baris "Hal"
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Hal" in t and "Penagihan" in t:
            hal_m = re.search(r'Penagihan Pembayaran (.+)', t)
            if hal_m:
                old_hal = hal_m.group(1).strip()
                new_hal = f"Jasa {project['jenis_pekerjaan']} {project['nama_klien_singkat']}"
                replacements[old_hal] = new_hal
                break

    return replacements


def generate_surat(template_bytes: bytes, project: dict, seq: int) -> bytes:
    """Inject project data into surat tagihan DOCX template.
    Supports both {{PLACEHOLDER}} template style and heuristic replace for filled templates.
    """
    doc = Document(io.BytesIO(template_bytes))

    tanggal = project["tanggal_tagihan_date"]
    fin = hitung_tagihan(project["imbalan_jasa_total"])
    nomor_sk = generate_nomor(tanggal, seq, "SK-OR", project["nama_klien_singkat"])
    tanggal_str = format_tanggal_indo(tanggal)
    total_fmt = f"Rp {fin['total']:,.0f}".replace(",", ".")

    # Detect mode
    mode = _detect_template_mode(template_bytes)

    if mode == "placeholder":
        # Mode 1: Template dengan {{PLACEHOLDER}} — exact replace
        # Pisah nomor menjadi prefix (YYMMDD.NNN) dan kode klien untuk template SRR aktual
        nomor_prefix  = f"{tanggal.strftime('%y%m%d')}.{seq:03d}"
        kode_pt       = project["nama_klien_singkat"].upper().replace(" ", "").replace(".", "")[:8]
        nomor_kwt_str = generate_nomor(tanggal, seq, "KWT.PJK", project["nama_klien_singkat"])

        tagih_ke      = project.get("tagih_ke", "Pertama")
        fee_fmt       = f"Rp {project['imbalan_jasa_total']:,.0f}".replace(",", ".")
        fee_tagih_fmt = f"{fin['imbalan_jasa']:,.0f}".replace(",", ".")
        dpp_fmt       = f"{fin['dpb_ppn']:,.0f}".replace(",", ".")
        ppn_fmt       = f"{fin['ppn']:,.0f}".replace(",", ".")
        jml_fmt       = f"{fin['total']:,.0f}".replace(",", ".")

        replacements = {
            # ── Template SRR aktual (token panjang didahulukan agar tidak partial-match)
            "{{Nomor_Srt}}//SRR-JK/KWT.PJK": nomor_kwt_str,
            "{{Nomor_Srt}}":        nomor_prefix,
            "{{Kode_PT}}":          kode_pt,
            "{{Tgl_Srt}}":          tanggal_str,
            "{{PEMBERI_TUGAS}}":    project["nama_klien"],
            "{{Pemberi_Tugas}}":    project["nama_klien"],
            "{{Alamat1}}":          project["alamat_baris1"],
            "{{Alamat2}}":          project["alamat_baris2"],
            "{{Kota}}":             project["kota"],
            "{{Kode_Pos}}":         project["kode_pos"],
            "{{Up}}":               project["up"],
            "{{Tagih_ke}}":         tagih_ke,
            "{{tagih_ke}}":         tagih_ke.lower(),
            "{{Pekerjaan}}":        project["jenis_pekerjaan"],
            "{{Nomor_Proposal}}":   project["nomor_proposal"],
            "{{Tanggal_Proposal}}": project["tanggal_proposal"],
            "{{Jumlah_Terbilang}}": fin["terbilang"],
            "{{Bank}}":             "Bank Mandiri KCP JKT Kalibata Rawajati",
            "{{Norek}}":            "126-0005748719",
            "{{title_Up}}":         project.get("title_up", "Bapak/Ibu"),
            "{{persentase}}":       project.get("persentase", "100%"),
            "{{Fee}}":              fee_fmt,
            "{{Fee_Tagih}}":        fee_tagih_fmt,
            "{{DPP}}":              dpp_fmt,
            "{{PPN}}":              ppn_fmt,
            "{{Jumlah}}":           jml_fmt,
            # ── Legacy placeholder (template default lama) — tetap didukung
            "{{NOMOR_SURAT}}":        nomor_sk,
            "{{TANGGAL_SURAT}}":      tanggal_str,
            "{{NAMA_KLIEN}}":         project["nama_klien"],
            "{{ALAMAT1}}":            project["alamat_baris1"],
            "{{ALAMAT2}}":            project["alamat_baris2"],
            "{{KOTA_POS}}":           f"{project['kota']} {project['kode_pos']}".strip(),
            "{{UP}}":                 project["up"],
            "{{JENIS_PEKERJAAN}}":    project["jenis_pekerjaan"],
            "{{TGL_PROPOSAL}}":       project["tanggal_proposal"],
            "{{TERBILANG}}":          fin["terbilang"],
            "{{TOTAL_ANGKA}}":        total_fmt,
            "{{RECEIVER}}":           project["receiver"],
            "{{NAMA_KLIEN_SINGKAT}}": project["nama_klien_singkat"],
        }
    else:
        # Mode 2: Template adalah surat yang sudah terisi — heuristic replace
        replacements = _build_heuristic_replacements(
            template_bytes, project, nomor_sk, tanggal_str, fin
        )
        # Tambahkan replace langsung untuk data yang paling kritis
        replacements.update({
            project.get("_template_nama_klien", "__NO_MATCH__"):  project["nama_klien"],
            project.get("_template_nomor_prop", "__NO_MATCH__"):  project["nomor_proposal"],
        })

    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        replace_in_table(table, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────
# TEMPLATE DOCX BUILDER (jika user tidak upload template)
# ─────────────────────────────────────────
def build_default_surat_template() -> bytes:
    """Build a basic surat tagihan DOCX template with placeholders."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        return p

    # Header KOP SURAT (simplified)
    add_para("SUWENDHO RINALDY DAN REKAN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para("KANTOR JASA PENILAI PUBLIK", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para("Nomor Izin Usaha KJPP: 2.09.0059 | Nomor Izin Cabang KJPP: 1138/KM.1/2017",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    add_para("Komplek Kalibata Indah Blok K16-17, Jl. Rawajati Timur, Pancoran, Jakarta Selatan 12750",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    add_para("─" * 80, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, space_after=12)

    # Nomor & Tanggal
    add_para("No. : {{Nomor_Srt}}/SRR-JK/SK-OR/{{Kode_PT}}\t\t{{Tgl_Srt}}", size=11)
    add_para("")

    # Kepada
    add_para("Kepada Yth.", size=11)
    add_para("{{PEMBERI_TUGAS}}", bold=True, size=11)
    add_para("{{Alamat1}}", bold=True, size=11)
    add_para("{{Alamat2}}", bold=True, size=11)
    add_para("{{Kota}} {{Kode_Pos}}", bold=True, size=11)
    add_para("")
    add_para("U.p.\t:\t{{Up}}", size=11)
    add_para("")

    # Hal
    add_para("Hal\t:\tPenagihan Pembayaran {{Tagih_ke}} {{Pekerjaan}}", bold=True, size=11)
    add_para("")
    add_para("Dengan hormat,", size=11)
    add_para("")

    # Body
    body = (
        "Menunjuk surat penawaran kami No. {{Nomor_Proposal}} tanggal {{Tanggal_Proposal}}, "
        "maka dengan ini kami mohon agar pembayaran {{tagih_ke}} untuk penugasan tersebut sebesar:"
    )
    add_para(body, size=11, space_after=12)
    add_para("")

    # Nominal box
    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_box.add_run("Rp {{Jumlah}}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    p_terb = doc.add_paragraph()
    p_terb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_terb.paragraph_format.space_after = Pt(12)
    r2 = p_terb.add_run("(  {{Jumlah_Terbilang}}  )")
    r2.font.size = Pt(11)
    r2.font.name = "Times New Roman"

    add_para("")
    add_para(
        "(kwitansi dan faktur PPN terlampir) dapat dibayarkan kepada kami dengan bilyet giro "
        "atau ditransfer ke rekening kami atas nama KJPP SUWENDHO RINALDY & REKAN di "
        "{{Bank}} dengan nomor rekening {{Norek}} pada kesempatan pertama.",
        size=11, space_after=12
    )
    add_para("")
    add_para("Demikianlah permohonan kami, atas perhatian dan kerja sama {{title_Up}} kami ucapkan terima kasih.", size=11)
    add_para("")
    add_para("Hormat kami,", size=11)
    add_para("")
    add_para("")
    add_para("{{RECEIVER}}", size=11)
    add_para("Rekan", size=11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────
# PREVIEW RENDERERS
# ─────────────────────────────────────────

def _fmt_rp(n: int) -> str:
    """Format angka sebagai Rp X.XXX.XXX"""
    return "Rp {:,.0f}".format(n).replace(",", ".")


def render_preview_surat(p: dict, seq: int) -> str:
    """
    Render surat tagihan sebagai HTML preview yang akurat.
    Menggunakan data live dari form (bukan template file).
    """
    tanggal = p.get("tanggal_tagihan_date", date.today())
    fin = hitung_tagihan(p["imbalan_jasa_total"])
    nomor_sk = generate_nomor(tanggal, seq, "SK-OR", p.get("nama_klien_singkat", "XXX"))
    tgl_str  = format_tanggal_indo(tanggal)
    total_fmt = _fmt_rp(fin["total"])

    # Baris alamat — hanya tampilkan baris yang ada isinya
    alamat_lines = []
    if p.get("alamat_baris1"): alamat_lines.append(p["alamat_baris1"])
    if p.get("alamat_baris2"): alamat_lines.append(p["alamat_baris2"])
    kota_pos = f"{p.get('kota','')} {p.get('kode_pos','')}".strip()
    if kota_pos: alamat_lines.append(kota_pos)
    alamat_html = "<br>".join(f"<strong>{a}</strong>" for a in alamat_lines)

    return f"""
<div style="
    font-family: 'Times New Roman', serif;
    font-size: 11.5px;
    line-height: 1.6;
    color: #1a1a1a;
    background: white;
    border: 1px solid #dde3ed;
    border-radius: 8px;
    padding: 24px 28px;
    max-height: 680px;
    overflow-y: auto;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
">
  <!-- KOP SURAT -->
  <div style="border-bottom: 2px solid #0f3460; padding-bottom: 10px; margin-bottom: 16px; display:flex; justify-content:space-between; align-items:flex-start;">
    <div>
      <div style="font-weight:800; font-size:13px; color:#0f3460; letter-spacing:0.3px;">SUWENDHO RINALDY DAN REKAN</div>
      <div style="font-size:9.5px; color:#555;">KANTOR JASA PENILAI PUBLIK</div>
      <div style="font-size:8.5px; color:#777;">No. Izin: 2.09.0059 | Cabang: 1138/KM.1/2017</div>
    </div>
    <div style="text-align:right; font-size:8.5px; color:#666; line-height:1.5;">
      Komplek Kalibata Indah Blok K16-17<br>
      Jl. Rawajati Timur, Pancoran<br>
      Jakarta Selatan 12750
    </div>
  </div>

  <!-- NOMOR & TANGGAL -->
  <table style="width:100%; margin-bottom:14px; font-size:11.5px;">
    <tr>
      <td style="width:60px; vertical-align:top; padding-right:4px;">No.</td>
      <td style="width:10px; vertical-align:top;">:</td>
      <td style="font-weight:600; color:#0f3460;">
        {nomor_sk if nomor_sk else '<span style="color:#aaa;">belum terisi</span>'}
      </td>
      <td style="text-align:right; color:#555;">{tgl_str}</td>
    </tr>
  </table>

  <!-- KEPADA -->
  <div style="margin-bottom:14px; font-size:11.5px;">
    <div>Kepada Yth.</div>
    <br>
    {alamat_html if alamat_lines else '<span style="color:#aaa; font-style:italic;">Alamat belum diisi</span>'}
    <br>
    <table style="margin-top:8px;">
      <tr>
        <td style="width:36px;">U.p.</td>
        <td style="width:10px;">:</td>
        <td><strong>{p.get("up","") or '<span style="color:#aaa;">—</span>'}</strong></td>
      </tr>
    </table>
  </div>

  <!-- HAL -->
  <table style="margin-bottom:16px; font-size:11.5px;">
    <tr>
      <td style="width:36px; vertical-align:top;">Hal</td>
      <td style="width:10px; vertical-align:top;">:</td>
      <td><strong>Penagihan Pembayaran Jasa {p.get("jenis_pekerjaan","") or "—"}
        {p.get("nama_klien_singkat","")}</strong></td>
    </tr>
  </table>

  <div style="margin-bottom:12px;">Dengan hormat,</div>

  <!-- BODY -->
  <div style="margin-bottom:16px; text-align:justify;">
    Menunjuk surat penawaran kami No.
    <strong>{p.get("nomor_proposal","") or '<span style="color:#aaa;">belum diisi</span>'}</strong>
    tanggal {p.get("tanggal_proposal","") or '<span style="color:#aaa;">—</span>'},
    maka dengan ini kami mohon agar pembayaran untuk penugasan tersebut sebesar:
  </div>

  <!-- NOMINAL -->
  <div style="
    text-align:center;
    font-size:15px;
    font-weight:700;
    color:#0f3460;
    padding: 10px 0 4px;
    letter-spacing:0.3px;
  ">{total_fmt if fin["total"] > 0 else '<span style="color:#aaa;">Rp —</span>'}</div>
  <div style="text-align:center; font-size:11px; color:#555; margin-bottom:14px;">
    ({fin["terbilang"] if fin["total"] > 0 else "—"})
  </div>

  <!-- RINCIAN MINI -->
  <table style="width:100%; font-size:10.5px; color:#555; margin-bottom:14px; border-top:1px solid #eee; padding-top:8px;">
    <tr>
      <td>Imbalan Jasa</td>
      <td style="text-align:right;">{_fmt_rp(fin["imbalan_jasa"])}</td>
    </tr>
    <tr>
      <td>DPP (× 11/12)</td>
      <td style="text-align:right;">{_fmt_rp(fin["dpb_ppn"])}</td>
    </tr>
    <tr>
      <td>PPN 12%</td>
      <td style="text-align:right;">{_fmt_rp(fin["ppn"])}</td>
    </tr>
    <tr style="font-weight:700; border-top:1px solid #ccc;">
      <td style="padding-top:4px;">Total</td>
      <td style="text-align:right; padding-top:4px; color:#0f3460;">{_fmt_rp(fin["total"])}</td>
    </tr>
  </table>

  <!-- FOOTER TEXT -->
  <div style="font-size:10.5px; color:#444; margin-bottom:20px; text-align:justify;">
    (kwitansi dan faktur PPN terlampir) dapat dibayarkan kepada kami dengan bilyet giro atau
    ditransfer ke rekening kami atas nama <strong>KJPP SUWENDHO RINALDY &amp; REKAN</strong>
    di Bank Mandiri KCP JKT Kalibata Rawajati No. Rek. <strong>126-0005748719</strong>
    pada kesempatan pertama.
  </div>

  <!-- TANDA TANGAN -->
  <div style="display:flex; justify-content:flex-end; font-size:11px; margin-top:8px;">
    <div style="text-align:center; width:160px;">
      <div>Hormat kami,</div>
      <div style="margin: 36px 0 4px; border-bottom:1px solid #333; width:120px; margin-left:auto; margin-right:auto;"></div>
      <div style="font-weight:700;">{p.get("receiver","Ocky Rinaldy")}</div>
      <div style="color:#666;">Rekan</div>
    </div>
  </div>
</div>
"""


def render_preview_kwitansi(p: dict, seq: int) -> str:
    """
    Render kwitansi sebagai HTML preview.
    """
    tanggal = p.get("tanggal_tagihan_date", date.today())
    fin = hitung_tagihan(p["imbalan_jasa_total"])
    nomor_kwt = generate_nomor(tanggal, seq, "KWT.PJK", p.get("nama_klien_singkat","XXX"))
    tgl_str   = format_tanggal_indo(tanggal)

    def row(label, val, bold_val=False):
        val_style = "font-weight:700;" if bold_val else ""
        color = "#0f3460" if bold_val else "#1a1a1a"
        return f"""
        <tr>
          <td style="padding:4px 8px 4px 0; color:#555; font-size:10px; white-space:nowrap;">{label}</td>
          <td style="padding:4px 0; text-align:right; font-size:10.5px; {val_style} color:{color};">{val}</td>
        </tr>"""

    return f"""
<div style="
    font-family: Arial, sans-serif;
    font-size: 11px;
    line-height: 1.6;
    color: #1a1a1a;
    background: white;
    border: 1px solid #dde3ed;
    border-radius: 8px;
    padding: 20px 24px;
    max-height: 680px;
    overflow-y: auto;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
">
  <!-- KOP -->
  <div style="border-bottom:2px solid #0f3460; padding-bottom:8px; margin-bottom:12px; text-align:center;">
    <div style="font-weight:800; font-size:12px; color:#0f3460;">SUWENDHO RINALDY DAN REKAN</div>
    <div style="font-size:8.5px; color:#555;">KANTOR JASA PENILAI PUBLIK</div>
  </div>

  <!-- TITLE + RECEIPT NO -->
  <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:14px;">
    <div style="font-size:18px; font-weight:800; color:#0f3460; letter-spacing:1px;">KWITANSI</div>
    <div style="text-align:right;">
      <div style="font-size:9px; color:#888;">Receipt No.</div>
      <div style="font-size:10px; font-weight:700; font-family:monospace; color:#0f3460;">
        {nomor_kwt}</div>
    </div>
  </div>

  <!-- DATA KLIEN -->
  <div style="background:#f8faff; border-radius:6px; padding:10px 12px; margin-bottom:14px; font-size:11px;">
    <div style="font-weight:700; font-size:11.5px; color:#0f3460; margin-bottom:2px;">
      {p.get("nama_klien","") or '<span style="color:#aaa; font-style:italic;">Nama klien belum diisi</span>'}
    </div>
    <div style="color:#555;">
      {p.get("alamat_baris1","")}{("<br>" + p.get("alamat_baris2","")) if p.get("alamat_baris2") else ""}
    </div>
    <div style="color:#666; font-size:10px;">
      {p.get("kota","")} {p.get("kode_pos","")}
    </div>
  </div>

  <!-- KETERANGAN -->
  <div style="font-size:10.5px; color:#444; margin-bottom:6px; border-left:3px solid #0f3460; padding-left:8px;">
    Pembayaran jasa <strong>{p.get("jenis_pekerjaan","—")}</strong>
    <span style="color:#777;"> · {p.get("nama_klien","")}</span><br>
    <span style="color:#888;">No. {p.get("nomor_proposal","—")} &nbsp;|&nbsp; {p.get("tanggal_proposal","—")}</span>
  </div>

  <!-- RINCIAN BIAYA -->
  <table style="width:100%; border-collapse:collapse; margin:12px 0 4px;">
    <thead>
      <tr style="background:#f1f5fe;">
        <th style="text-align:left; padding:5px 8px 5px 0; font-size:9.5px; color:#666; font-weight:600;">URAIAN</th>
        <th style="text-align:right; padding:5px 0; font-size:9.5px; color:#666; font-weight:600;">JUMLAH (Rp)</th>
      </tr>
    </thead>
    <tbody>
      {row("Imbalan Jasa", _fmt_rp(fin["imbalan_jasa"]))}
      {row("Dasar Pengenaan PPN (× 11/12)", _fmt_rp(fin["dpb_ppn"]))}
      {row("PPN 12%", _fmt_rp(fin["ppn"]))}
    </tbody>
    <tfoot>
      <tr style="border-top:2px solid #0f3460;">
        <td style="padding:6px 8px 2px 0; font-weight:700; font-size:11px;">TOTAL</td>
        <td style="text-align:right; font-weight:800; font-size:13px; color:#0f3460; padding-top:6px;">
          {_fmt_rp(fin["total"]) if fin["total"] > 0 else "—"}
        </td>
      </tr>
    </tfoot>
  </table>

  <!-- TERBILANG -->
  <div style="
    background:#fffbeb;
    border:1px dashed #d4a017;
    border-radius:5px;
    padding:8px 12px;
    font-size:10.5px;
    color:#6b4c00;
    margin:10px 0 14px;
    font-style:italic;
  ">
    ( {fin["terbilang"] if fin["total"] > 0 else "—"} )
  </div>

  <!-- FOOTER -->
  <div style="display:flex; justify-content:space-between; align-items:flex-end; font-size:10px; color:#666; margin-top:8px;">
    <div style="font-size:9px; color:#888; max-width:60%;">
      Pembayaran melalui Bank Mandiri KCP JKT Kalibata Rawajati<br>
      a.n. KJPP Suwendho Rinaldy &amp; Rekan · No. Rek. 126-0005748719
    </div>
    <div style="text-align:center; min-width:120px;">
      <div>Jakarta, {tgl_str}</div>
      <div style="margin:28px 0 4px; border-bottom:1px solid #555; width:100px; margin-left:auto; margin-right:auto;"></div>
      <div style="font-weight:700;">{p.get("receiver","Ocky Rinaldy")}</div>
      <div style="color:#888;">Rekan</div>
    </div>
  </div>
</div>
"""


def render_proposal_text(p: dict) -> str:
    """
    Tampilkan teks mentah proposal yang tersimpan.
    Highlight field-field yang berhasil diekstrak.
    """
    proposal_bytes = p.get("_proposal_bytes")
    proposal_name  = p.get("_proposal_name", "")
    if not proposal_bytes:
        return None

    # Ekstrak teks (hanya DOCX; PDF langsung dikirim ke Claude API)
    try:
        if proposal_name.lower().endswith(".pdf"):
            import html as html_mod
            return f"""
<div style="
    font-family: sans-serif; font-size: 11px; color: #6b7280;
    background: #f9fafb; border: 1px solid #e5e7eb;
    border-radius: 8px; padding: 16px; text-align: center;
">
    📄 <strong>{html_mod.escape(proposal_name)}</strong><br><br>
    Preview teks tidak tersedia untuk file PDF.<br>
    Data diekstrak langsung oleh AI dari file PDF asli.
</div>"""
        doc = Document(io.BytesIO(proposal_bytes))
        raw = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception:
        return None

    if not raw:
        return None

    # Highlight nilai-nilai yang diekstrak dalam teks proposal
    import html as html_mod
    safe = html_mod.escape(raw[:5000])  # batasi agar tidak terlalu panjang

    # Highlight setiap nilai yang cocok (case-insensitive)
    highlights = [
        p.get("nama_klien",""), p.get("nomor_proposal",""),
        p.get("tanggal_proposal",""), p.get("nama_klien_singkat",""),
        p.get("alamat_baris1",""), p.get("up",""),
    ]
    for val in highlights:
        if val and len(val) > 3:
            safe_val = html_mod.escape(val)
            safe = safe.replace(safe_val,
                f'<mark style="background:#fef08a; border-radius:3px; padding:0 2px;">{safe_val}</mark>')

    return f"""
<div style="
    font-family: 'JetBrains Mono', 'Courier New', monospace;
    font-size: 10px;
    line-height: 1.7;
    color: #374151;
    background: #f9fafb;
    border: 1px solid #e5e7eb;
    border-radius: 8px;
    padding: 16px;
    max-height: 680px;
    overflow-y: auto;
    white-space: pre-wrap;
    word-break: break-word;
">
<div style="font-family:sans-serif; font-size:10px; color:#9ca3af; margin-bottom:10px; padding-bottom:6px; border-bottom:1px solid #e5e7eb;">
  📄 {html_mod.escape(proposal_name)} · {len(raw):,} karakter · (5.000 karakter pertama)
</div>{safe}{"..." if len(raw) > 5000 else ""}
</div>
"""


# ─────────────────────────────────────────
# SESSION STATE INIT
# ─────────────────────────────────────────
if "projects"    not in st.session_state: st.session_state.projects = []
if "sk_template" not in st.session_state: st.session_state.sk_template = None
if "seq_counter" not in st.session_state: st.session_state.seq_counter = 1
if "edit_idx"    not in st.session_state: st.session_state.edit_idx = None


# ─────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────
st.markdown("""
<div class="srr-header">
  <div>
    <h1>📄 SRR Billing Generator</h1>
    <p>Generator Kwitansi & Surat Penagihan — KJPP Suwendho Rinaldy dan Rekan</p>
  </div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# SIDEBAR — SETTINGS + TEMPLATES
# ─────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Pengaturan")

    receiver = st.text_input("Nama Penandatangan", value="Ocky Rinaldy")
    seq_start = st.number_input("Nomor Urut Awal", min_value=1, max_value=999,
                                 value=st.session_state.seq_counter)
    if seq_start != st.session_state.seq_counter:
        st.session_state.seq_counter = seq_start

    st.divider()
    st.markdown("### 📎 Upload Template")

    sk_upload = st.file_uploader(
        "Template Dokumen (.docx)",
        type=["docx"],
        key="sk_up",
        help="Upload template DOCX yang berisi placeholder {{Nomor_Srt}}, {{PEMBERI_TUGAS}}, dll."
    )
    if sk_upload:
        st.session_state.sk_template = sk_upload.read()
        mode = _detect_template_mode(st.session_state.sk_template)
        if mode == "placeholder":
            st.success("✅ Template tersimpan — mode **{{Placeholder}}** terdeteksi")
        else:
            st.warning(
                "⚠️ Template tersimpan — **mode Heuristic**. "
                "Untuk hasil terbaik, gunakan placeholder `{{Nomor_Srt}}`, `{{PEMBERI_TUGAS}}`, dll."
            )

    if not st.session_state.sk_template:
        st.info("ℹ️ Belum ada template — akan gunakan template default SRR")

    st.divider()
    st.markdown("#### 💡 Placeholder Surat")
    with st.expander("Lihat daftar placeholder"):
        st.code("""{{Nomor_Srt}}        ← prefix nomor (YYMMDD.NNN)
{{Kode_PT}}         ← kode klien (misal: PTRO)
{{Tgl_Srt}}         ← tanggal surat
{{PEMBERI_TUGAS}}   ← nama klien lengkap
{{Alamat1}}         ← alamat baris 1
{{Alamat2}}         ← alamat baris 2
{{Kota}}            ← kota
{{Kode_Pos}}        ← kode pos
{{Up}}              ← jabatan penerima
{{Tagih_ke}}        ← urutan penagihan (Pertama/Kedua/…)
{{tagih_ke}}        ← idem, huruf kecil
{{Pekerjaan}}       ← jenis pekerjaan
{{Nomor_Proposal}}  ← nomor proposal
{{Tanggal_Proposal}}← tanggal proposal
{{Jumlah_Terbilang}}← terbilang total
{{Fee}}             ← imbalan jasa (Rp)
{{Fee_Tagih}}       ← imbalan jasa yang ditagih
{{DPP}}             ← dasar pengenaan PPN
{{PPN}}             ← nilai PPN 12%
{{Jumlah}}          ← total (Fee + PPN)
{{Bank}}            ← nama bank
{{Norek}}           ← nomor rekening
{{title_Up}}        ← sapaan penutup (Bapak/Ibu)
{{persentase}}      ← persentase tagih (100%)
{{RECEIVER}}        ← nama penandatangan""", language="text")
        st.caption("Template lama ({{NOMOR_SURAT}}, {{NAMA_KLIEN}}, dll.) tetap didukung.")


# ─────────────────────────────────────────
# MAIN TABS
# ─────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["📥 Input Proyek", "📋 Review & Edit", "📦 Download"])

# ════════════════════════════════════════
# TAB 1 — INPUT
# ════════════════════════════════════════
with tab1:
    st.markdown('<div class="step-pill">LANGKAH 1 — Tambah Proyek</div>', unsafe_allow_html=True)

    input_mode = st.radio(
        "Mode input:",
        ["🤖 Upload Proposal (Ekstraksi AI)", "✏️ Input Manual"],
        horizontal=True
    )

    # ── AI MODE
    if "🤖" in input_mode:
        uploaded_proposals = st.file_uploader(
            "Upload Proposal / Kontrak (PDF atau DOCX)",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            help="Bisa upload beberapa file sekaligus untuk batch processing"
        )

        if uploaded_proposals:
            if st.button("🤖 Ekstrak Data dengan AI", type="primary",
                         disabled=len(uploaded_proposals) == 0):
                progress = st.progress(0)
                for i, f in enumerate(uploaded_proposals):
                    with st.spinner(f"Membaca {f.name}..."):
                        try:
                            file_bytes = f.read()
                            data = extract_proposal_data(file_bytes, f.name)
                            data["receiver"] = receiver
                            data["_source_file"] = f.name
                            # Simpan bytes proposal untuk preview
                            data["_proposal_bytes"] = file_bytes
                            data["_proposal_name"] = f.name
                            data["tanggal_tagihan_date"] = date.today()
                            st.session_state.projects.append(data)
                        except Exception as e:
                            st.error(f"Gagal proses {f.name}: {e}")
                    progress.progress((i + 1) / len(uploaded_proposals))

                st.success(f"✅ {len(uploaded_proposals)} proyek berhasil diekstrak. Lanjut ke tab **Review & Edit** untuk memeriksa & download.")
                st.balloons()

    # ── MANUAL MODE
    else:
        with st.form("manual_form", clear_on_submit=True):
            st.markdown("#### Data Klien")
            c1, c2 = st.columns(2)
            with c1:
                m_nama       = st.text_input("Nama Klien *")
                m_singkat    = st.text_input("Kode/Singkatan Klien *", help="Contoh: PTRO, IDXSTI")
                m_alamat1    = st.text_input("Alamat Baris 1")
                m_alamat2    = st.text_input("Alamat Baris 2")
            with c2:
                m_kota       = st.text_input("Kota")
                m_kodepos    = st.text_input("Kode Pos")
                m_up         = st.text_input("U.p. (Jabatan)", value="Direksi")
                m_receiver_f = st.text_input("Nama Penandatangan", value=receiver)

            st.markdown("#### Detail Penugasan")
            c3, c4 = st.columns(2)
            with c3:
                m_pekerjaan  = st.text_input("Jenis Pekerjaan *",
                                              help="Contoh: Penilaian Saham dan Pendapat Kewajaran")
                m_noprop     = st.text_input("Nomor Proposal")
                m_tglprop    = st.text_input("Tanggal Proposal", help="Contoh: 12 Januari 2026")
                m_tagih_ke   = st.text_input("Tagih ke (Urutan Penagihan)", value="Pertama",
                                              help="Contoh: Pertama, Kedua, Ketiga")
            with c4:
                m_fee        = st.number_input("Imbalan Jasa (Rp, sebelum PPN) *",
                                               min_value=0, step=1000000)
                m_tgltagih   = st.date_input("Tanggal Tagihan", value=date.today())

            # Preview perhitungan
            if m_fee > 0:
                fin_prev = hitung_tagihan(m_fee)
                st.info(
                    f"**Preview:** DPP = Rp {fin_prev['dpb_ppn']:,.0f} | "
                    f"PPN 12% = Rp {fin_prev['ppn']:,.0f} | "
                    f"**Total = Rp {fin_prev['total']:,.0f}**"
                )

            submitted = st.form_submit_button("➕ Tambah ke Daftar Proyek", type="primary")
            if submitted:
                if not m_nama or not m_singkat or not m_pekerjaan or m_fee <= 0:
                    st.error("⚠️ Isi semua field wajib (*)")
                else:
                    proj = {
                        "nama_klien": m_nama,
                        "nama_klien_singkat": m_singkat,
                        "alamat_baris1": m_alamat1,
                        "alamat_baris2": m_alamat2,
                        "kota": m_kota,
                        "kode_pos": m_kodepos,
                        "up": m_up,
                        "jenis_pekerjaan": m_pekerjaan,
                        "nomor_proposal": m_noprop,
                        "tanggal_proposal": m_tglprop,
                        "tagih_ke": m_tagih_ke,
                        "imbalan_jasa_total": int(m_fee),
                        "tanggal_tagihan_date": m_tgltagih,
                        "receiver": m_receiver_f,
                        "_source_file": "Manual",
                    }
                    st.session_state.projects.append(proj)
                    st.success(f"✅ {m_nama} ditambahkan!")

    # Show current count
    if st.session_state.projects:
        st.markdown(f"---\n**{len(st.session_state.projects)} proyek** dalam antrian. Lanjut ke tab **Review & Edit**.")


# ════════════════════════════════════════
# TAB 2 — REVIEW & EDIT
# ════════════════════════════════════════
with tab2:
    st.markdown('<div class="step-pill">LANGKAH 2 — Review & Edit Data</div>', unsafe_allow_html=True)

    if not st.session_state.projects:
        st.info("Belum ada proyek. Tambahkan di tab **Input Proyek** terlebih dahulu.")
    else:
        # ── Summary table
        summary_rows = []
        for i, p in enumerate(st.session_state.projects):
            fin = hitung_tagihan(p["imbalan_jasa_total"])
            summary_rows.append({
                "#": i + 1,
                "Klien": p["nama_klien"],
                "Kode": p["nama_klien_singkat"],
                "Pekerjaan": p["jenis_pekerjaan"][:45] + "…" if len(p["jenis_pekerjaan"]) > 45 else p["jenis_pekerjaan"],
                "Imbalan Jasa": f"Rp {p['imbalan_jasa_total']:,.0f}",
                "Total + PPN": f"Rp {fin['total']:,.0f}",
                "Sumber": p.get("_source_file", "-"),
            })

        import pandas as pd
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

        st.divider()

        # ── Pilih proyek
        edit_idx = st.selectbox(
            "Pilih proyek untuk diedit:",
            options=list(range(len(st.session_state.projects))),
            format_func=lambda i: f"#{i+1} — {st.session_state.projects[i]['nama_klien']}"
        )

        p = st.session_state.projects[edit_idx]
        seq_preview = st.session_state.seq_counter + edit_idx

        # ── LAYOUT UTAMA: Form (kiri) | Preview tabs (kanan)
        col_form, col_prev = st.columns([5, 6], gap="large")

        # ───────────────────────────────
        # KOLOM KIRI — Edit Form
        # ───────────────────────────────
        with col_form:
            st.markdown("### ✏️ Edit Detail")

            with st.form(f"edit_form_{edit_idx}"):
                st.markdown("**Data Klien**")
                e_nama    = st.text_input("Nama Klien", value=p["nama_klien"])
                e_singkat = st.text_input("Kode Klien", value=p["nama_klien_singkat"],
                                          help="2-8 huruf, contoh: PTRO, IDXSTI")
                e_alamat1 = st.text_input("Alamat 1 (Gedung/Wisma)", value=p["alamat_baris1"])
                e_alamat2 = st.text_input("Alamat 2 (Jalan)", value=p["alamat_baris2"])

                ca, cb = st.columns(2)
                with ca:
                    e_kota    = st.text_input("Kota", value=p["kota"])
                with cb:
                    e_kodepos = st.text_input("Kode Pos", value=str(p["kode_pos"]))

                e_up = st.text_input("U.p. (Jabatan penerima)", value=p["up"])

                st.markdown("**Detail Penugasan**")
                e_pekerjaan = st.text_input("Jenis Pekerjaan", value=p["jenis_pekerjaan"])
                e_noprop    = st.text_input("Nomor Proposal", value=p["nomor_proposal"])
                e_tglprop   = st.text_input("Tanggal Proposal", value=p["tanggal_proposal"],
                                            help="Contoh: 12 Januari 2026")
                e_tagih_ke  = st.text_input("Tagih ke (Urutan Penagihan)",
                                            value=p.get("tagih_ke", "Pertama"),
                                            help="Contoh: Pertama, Kedua, Ketiga")

                st.markdown("**Tagihan**")
                e_fee = st.number_input("Imbalan Jasa (Rp, sebelum PPN)",
                                        value=int(p["imbalan_jasa_total"]),
                                        min_value=0, step=1_000_000,
                                        help="Total fee sebelum PPN. DPP & PPN dihitung otomatis.")
                e_tgltagih = st.date_input("Tanggal Tagihan",
                                           value=p.get("tanggal_tagihan_date", date.today()))
                e_receiver = st.text_input("Penandatangan", value=p["receiver"])

                # Preview angka live
                if e_fee > 0:
                    fin_e = hitung_tagihan(e_fee)
                    st.info(
                        f"DPP = {_fmt_rp(fin_e['dpb_ppn'])} · "
                        f"PPN = {_fmt_rp(fin_e['ppn'])} · "
                        f"**Total = {_fmt_rp(fin_e['total'])}**"
                    )

                cs, cd = st.columns([3, 1])
                with cs:
                    saved = st.form_submit_button("💾 Simpan Perubahan", type="primary",
                                                   use_container_width=True)
                with cd:
                    deleted = st.form_submit_button("🗑️ Hapus", use_container_width=True)

                if saved:
                    st.session_state.projects[edit_idx].update({
                        "nama_klien":          e_nama,
                        "nama_klien_singkat":  e_singkat,
                        "alamat_baris1":       e_alamat1,
                        "alamat_baris2":       e_alamat2,
                        "kota":                e_kota,
                        "kode_pos":            e_kodepos,
                        "up":                  e_up,
                        "jenis_pekerjaan":     e_pekerjaan,
                        "nomor_proposal":      e_noprop,
                        "tanggal_proposal":    e_tglprop,
                        "tagih_ke":            e_tagih_ke,
                        "imbalan_jasa_total":  int(e_fee),
                        "tanggal_tagihan_date":e_tgltagih,
                        "receiver":            e_receiver,
                    })
                    st.success("✅ Data diperbarui!")
                    st.rerun()

                if deleted:
                    st.session_state.projects.pop(edit_idx)
                    st.success("Proyek dihapus.")
                    st.rerun()

        # ───────────────────────────────
        # KOLOM KANAN — Preview Panel
        # ───────────────────────────────
        with col_prev:
            st.markdown("### 👁️ Preview Dokumen")

            # Build preview project dari nilai form SAAT INI (sebelum disimpan)
            # Gunakan data tersimpan (p) karena form belum di-submit
            p_preview = dict(p)

            has_proposal = bool(p.get("_proposal_bytes"))

            if has_proposal:
                prev_tab_surat, prev_tab_kwt, prev_tab_proposal = st.tabs([
                    "📝 Surat Tagihan", "📊 Kwitansi", "📄 Teks Proposal"
                ])
            else:
                prev_tab_surat, prev_tab_kwt = st.tabs([
                    "📝 Surat Tagihan", "📊 Kwitansi"
                ])

            with prev_tab_surat:
                if p_preview["imbalan_jasa_total"] == 0:
                    st.warning("⚠️ Imbalan Jasa belum diisi — preview nominal kosong.")
                surat_html = render_preview_surat(p_preview, seq_preview)
                st.markdown(surat_html, unsafe_allow_html=True)

                # Checklist validasi surat
                st.markdown("---")
                st.markdown("**Checklist Validasi**")
                checks_surat = {
                    "Nama klien":       bool(p_preview.get("nama_klien")),
                    "Alamat":           bool(p_preview.get("alamat_baris1")),
                    "Nomor proposal":   bool(p_preview.get("nomor_proposal")),
                    "Tanggal proposal": bool(p_preview.get("tanggal_proposal")),
                    "Jenis pekerjaan":  bool(p_preview.get("jenis_pekerjaan")),
                    "Imbalan jasa > 0": p_preview.get("imbalan_jasa_total", 0) > 0,
                    "Kode klien":       bool(p_preview.get("nama_klien_singkat")),
                }
                cols_check = st.columns(2)
                for ci, (label, ok) in enumerate(checks_surat.items()):
                    with cols_check[ci % 2]:
                        icon = "✅" if ok else "❌"
                        color = "#166534" if ok else "#991b1b"
                        st.markdown(
                            f'<span style="font-size:12px; color:{color};">{icon} {label}</span>',
                            unsafe_allow_html=True
                        )

            with prev_tab_kwt:
                if p_preview["imbalan_jasa_total"] == 0:
                    st.warning("⚠️ Imbalan Jasa belum diisi — preview nominal kosong.")
                kwt_html = render_preview_kwitansi(p_preview, seq_preview)
                st.markdown(kwt_html, unsafe_allow_html=True)

                # Checklist validasi kwitansi
                st.markdown("---")
                st.markdown("**Checklist Validasi**")
                checks_kwt = {
                    "Nama klien":     bool(p_preview.get("nama_klien")),
                    "Kode klien":     bool(p_preview.get("nama_klien_singkat")),
                    "Alamat":         bool(p_preview.get("alamat_baris1")),
                    "Kota":           bool(p_preview.get("kota")),
                    "Kode pos":       bool(p_preview.get("kode_pos")),
                    "Fee > 0":        p_preview.get("imbalan_jasa_total", 0) > 0,
                    "No. proposal":   bool(p_preview.get("nomor_proposal")),
                    "Tgl. proposal":  bool(p_preview.get("tanggal_proposal")),
                }
                cols_ck2 = st.columns(2)
                for ci, (label, ok) in enumerate(checks_kwt.items()):
                    with cols_ck2[ci % 2]:
                        icon = "✅" if ok else "❌"
                        color = "#166534" if ok else "#991b1b"
                        st.markdown(
                            f'<span style="font-size:12px; color:{color};">{icon} {label}</span>',
                            unsafe_allow_html=True
                        )

            if has_proposal:
                with prev_tab_proposal:
                    st.caption(
                        "Teks asli proposal yang diupload. "
                        "**Kuning** = nilai yang berhasil diekstrak AI. "
                        "Jika ada data yang tidak cocok, edit di panel kiri lalu Simpan."
                    )
                    proposal_html = render_proposal_text(p_preview)
                    if proposal_html:
                        st.markdown(proposal_html, unsafe_allow_html=True)
                    else:
                        st.info("Tidak bisa membaca teks proposal.")

                    # Tabel perbandingan: Ekstraksi AI vs Data saat ini
                    st.markdown("---")
                    st.markdown("**Perbandingan: Hasil Ekstraksi AI vs Data Tersimpan**")
                    compare_fields = [
                        ("Nama Klien",       p_preview.get("nama_klien","")),
                        ("Kode Klien",       p_preview.get("nama_klien_singkat","")),
                        ("Alamat 1",         p_preview.get("alamat_baris1","")),
                        ("Alamat 2",         p_preview.get("alamat_baris2","")),
                        ("Kota",             p_preview.get("kota","")),
                        ("Kode Pos",         str(p_preview.get("kode_pos",""))),
                        ("Nomor Proposal",   p_preview.get("nomor_proposal","")),
                        ("Tanggal Proposal", p_preview.get("tanggal_proposal","")),
                        ("Imbalan Jasa",     _fmt_rp(p_preview.get("imbalan_jasa_total",0))),
                        ("Jenis Pekerjaan",  p_preview.get("jenis_pekerjaan","")),
                    ]
                    df_compare = pd.DataFrame(compare_fields, columns=["Field", "Nilai Tersimpan"])
                    st.dataframe(df_compare, use_container_width=True, hide_index=True)


# ════════════════════════════════════════
# TAB 3 — DOWNLOAD
# ════════════════════════════════════════
with tab3:
    st.markdown('<div class="step-pill">LANGKAH 3 — Generate & Download</div>', unsafe_allow_html=True)

    if not st.session_state.projects:
        st.info("Belum ada proyek. Tambahkan terlebih dahulu di tab **Input Proyek**.")
    else:
        n = len(st.session_state.projects)
        st.success(f"**{n} proyek** siap di-generate.")

        doc_tpl = st.session_state.sk_template or build_default_surat_template()

        if st.button("⚡ Generate Semua Dokumen", type="primary", use_container_width=True):
            zip_buf = io.BytesIO()
            errors  = []

            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, proj in enumerate(st.session_state.projects):
                    seq     = st.session_state.seq_counter + i
                    kode    = proj["nama_klien_singkat"].upper()[:8]
                    tanggal = proj.get("tanggal_tagihan_date", date.today())
                    prefix  = tanggal.strftime("%y%m%d")

                    try:
                        doc_bytes = generate_surat(doc_tpl, proj, seq)
                        doc_fname = f"{prefix}_{seq:03d}-SK-OR-{kode}_Tagihan.docx"
                        zf.writestr(doc_fname, doc_bytes)
                    except Exception as e:
                        errors.append(f"Dokumen #{i+1} ({kode}): {e}")

            zip_buf.seek(0)

            if errors:
                for err in errors:
                    st.error(f"⚠️ {err}")

            st.download_button(
                label=f"📦 Download ZIP ({n} dokumen tagihan)",
                data=zip_buf.read(),
                file_name=f"SRR_Tagihan_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                mime="application/zip",
                use_container_width=True,
                type="primary",
            )

        st.divider()
        st.markdown("#### 📄 Download Individual")

        doc_tpl_ind = st.session_state.sk_template or build_default_surat_template()

        for i, proj in enumerate(st.session_state.projects):
            seq     = st.session_state.seq_counter + i
            kode    = proj["nama_klien_singkat"].upper()[:8]
            tanggal = proj.get("tanggal_tagihan_date", date.today())
            fin     = hitung_tagihan(proj["imbalan_jasa_total"])

            with st.expander(f"#{i+1} — {proj['nama_klien']} | Total: Rp {fin['total']:,.0f}"):
                try:
                    doc_b = generate_surat(doc_tpl_ind, proj, seq)
                    st.download_button(
                        f"📝 Download Dokumen Tagihan ({kode})",
                        data=doc_b,
                        file_name=f"SKOR_{kode}_{tanggal.strftime('%y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"doc_{i}",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Error generate dokumen: {e}")

        st.divider()
        if st.button("🗑️ Hapus Semua Proyek", type="secondary"):
            st.session_state.projects = []
            st.session_state.seq_counter = 1
            st.rerun()
